
"""
Generador de reportes en Word y Excel
"""
import os
from datetime import datetime
import logging
from io import BytesIO
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import numpy as np

logger = logging.getLogger(__name__)


def create_probability_graphs(predictions):
    """Crea gráficos de distribución de probabilidades para cada modelo"""
    class_names = ["Células Superficiales", "Células Intermedias", 
                   "Células Parabasales", "Células Koilocíticas", 
                   "Células Displásicas"]
    colors_list = ['#FF6B6B', '#4ECDC4', '#FFE66D', '#95E1D3', '#F38181']
    images = []
    
    for model_name, pred in predictions.items():
        probabilities = pred.get('probabilities', [0]*5)
        
        fig, ax = plt.subplots(figsize=(6, 4))
        bars = ax.bar(class_names, probabilities, color=colors_list)
        ax.set_title(f'{model_name} - Distribución de Probabilidades')
        ax.set_ylabel('Probabilidad')
        ax.set_ylim([0, 1])
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)
        images.append(buf)
        plt.close()
    
    return images


def create_consensus_graph(predictions):
    """Crea un gráfico de pastel para el consenso entre modelos"""
    class_counts = {}
    total_models = len(predictions)
    
    for pred in predictions.values():
        cls = pred.get('class_friendly', pred.get('class', 'N/A'))
        class_counts[cls] = class_counts.get(cls, 0) + 1
    
    fig, ax = plt.subplots(figsize=(5, 5))
    colors_list = ['#00D25B', '#FF6B6B', '#4ECDC4', '#FFE66D', '#95E1D3']
    ax.pie(class_counts.values(), labels=class_counts.keys(), 
           autopct='%1.1f%%', colors=colors_list[:len(class_counts)],
           startangle=90)
    ax.set_title('Consenso entre Modelos')
    
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=100)
    buf.seek(0)
    plt.close()
    return buf


def generate_word_report(predictions, image_info, patient_info):
    """
    Genera un reporte en Word (.docx)
    """
    try:
        doc = Document()
        
        # Título
        title = doc.add_heading('Reporte de Análisis de Células Cervicales', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitle = doc.add_heading('Sistema SiPakMed-AI', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha y datos del paciente
        doc.add_paragraph(f'Fecha y Hora del Análisis: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        if patient_info and patient_info.get('name'):
            doc.add_paragraph(f'Paciente: {patient_info["name"]}')
        if patient_info and patient_info.get('id'):
            doc.add_paragraph(f'ID de Paciente: {patient_info["id"]}')
        
        # Sección de Análisis de Imagen
        doc.add_heading('1. Análisis de la Imagen', level=2)
        if isinstance(image_info, dict):
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'
            table.rows[0].cells[0].text = 'Nombre del Archivo'
            table.rows[0].cells[1].text = str(image_info.get('filename', 'N/A'))
            table.rows[1].cells[0].text = 'Dimensiones'
            table.rows[1].cells[1].text = str(image_info.get('size', 'N/A'))
            table.rows[2].cells[0].text = 'Formato'
            table.rows[2].cells[1].text = str(image_info.get('format', 'N/A'))
            table.rows[3].cells[0].text = 'Modo'
            table.rows[3].cells[1].text = str(image_info.get('mode', 'N/A'))
        
        # Sección de Resultados
        doc.add_heading('2. Resultados de Predicción', level=2)
        
        # Tabla de predicciones por modelo
        if predictions:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Modelo'
            hdr_cells[1].text = 'Clase Predicha'
            hdr_cells[2].text = 'Confianza (%)'
            hdr_cells[3].text = 'Tipo de Modelo'
            
            for model_name, pred in predictions.items():
                row_cells = table.add_row().cells
                class_name = pred.get('class_friendly', pred.get('predicted_class', 'N/A'))
                confidence = pred.get('confidence', 0)
                model_type = "Híbrido" if any(h in model_name for h in ['Hybrid', 'Ensemble', 'Multiscale']) else "Clásico"
                
                row_cells[0].text = model_name
                row_cells[1].text = str(class_name)
                row_cells[2].text = f"{confidence:.2f}"
                row_cells[3].text = model_type
            
            # Resumen
            doc.add_heading('Resumen del Análisis', level=3)
            total_models = len(predictions)
            avg_confidence = sum(p.get('confidence', 0) for p in predictions.values()) / total_models
            class_counts = {}
            for pred in predictions.values():
                cls = pred.get('class_friendly', pred.get('predicted_class', 'N/A'))
                class_counts[cls] = class_counts.get(cls, 0) + 1
            most_common = max(class_counts.items(), key=lambda x: x[1]) if class_counts else ("N/A", 0)
            
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            summary_table.rows[0].cells[0].text = 'Total de Modelos'
            summary_table.rows[0].cells[1].text = str(total_models)
            summary_table.rows[1].cells[0].text = 'Confianza Promedio'
            summary_table.rows[1].cells[1].text = f"{avg_confidence:.2f}%"
            summary_table.rows[2].cells[0].text = 'Clase con Mayor Consenso'
            summary_table.rows[2].cells[1].text = f"{most_common[0]}"
            summary_table.rows[3].cells[0].text = 'Votos para la Clase Principal'
            summary_table.rows[3].cells[1].text = f"{most_common[1]} de {total_models}"
        
        # Sección de Interpretación
        doc.add_heading('3. Interpretación y Información del Sistema', level=2)
        system_table = doc.add_table(rows=4, cols=2)
        system_table.style = 'Light Grid Accent 1'
        system_table.rows[0].cells[0].text = 'Modelos Disponibles'
        system_table.rows[0].cells[1].text = '5 (3 clásicos + 2 híbridos)'
        system_table.rows[1].cells[0].text = 'Precisión del Sistema'
        system_table.rows[1].cells[1].text = '84-93%'
        system_table.rows[2].cells[0].text = 'Clases de Células'
        system_table.rows[2].cells[1].text = '5'
        system_table.rows[3].cells[0].text = 'Modo de Procesamiento'
        system_table.rows[3].cells[1].text = 'CPU/GPU'
        
        # Guardar el documento en un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        logger.error(f"Error en generate_word_report: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return None


def generate_excel_report(predictions, image_info, patient_info):
    """
    Genera un reporte en Excel (.xlsx)
    """
    try:
        buffer = BytesIO()
        
        # Crear un escritor de Excel
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            # Hoja 1: Resumen
            summary_data = {
                'Campo': ['Fecha y Hora del Análisis', 
                          'Nombre del Paciente' if (patient_info and patient_info.get('name')) else '',
                          'ID del Paciente' if (patient_info and patient_info.get('id')) else '',
                          'Nombre del Archivo',
                          'Dimensiones',
                          'Formato',
                          'Modo',
                          'Total de Modelos',
                          'Clases de Células',
                          'Precisión del Sistema'],
                'Valor': [datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                          patient_info['name'] if (patient_info and patient_info.get('name')) else '',
                          patient_info['id'] if (patient_info and patient_info.get('id')) else '',
                          str(image_info.get('filename', 'N/A')) if isinstance(image_info, dict) else 'N/A',
                          str(image_info.get('size', 'N/A')) if isinstance(image_info, dict) else 'N/A',
                          str(image_info.get('format', 'N/A')) if isinstance(image_info, dict) else 'N/A',
                          str(image_info.get('mode', 'N/A')) if isinstance(image_info, dict) else 'N/A',
                          len(predictions),
                          '5',
                          '84-93%']
            }
            df_summary = pd.DataFrame(summary_data)
            df_summary = df_summary[df_summary['Campo'] != '']
            df_summary.to_excel(writer, sheet_name='Resumen', index=False)
            
            # Hoja 2: Predicciones por Modelo
            if predictions:
                pred_list = []
                for model_name, pred in predictions.items():
                    class_name = pred.get('class_friendly', pred.get('predicted_class', 'N/A'))
                    confidence = pred.get('confidence', 0)
                    model_type = "Híbrido" if any(h in model_name for h in ['Hybrid', 'Ensemble', 'Multiscale']) else "Clásico"
                    pred_list.append({
                        'Modelo': model_name,
                        'Clase Predicha': str(class_name),
                        'Confianza (%)': round(confidence, 2),
                        'Tipo de Modelo': model_type,
                        'Probabilidades': ', '.join([f'{p:.2f}' for p in pred.get('probabilities', [])])
                    })
                
                df_pred = pd.DataFrame(pred_list)
                df_pred.to_excel(writer, sheet_name='Predicciones', index=False)
            
            # Hoja 3: Consenso
            if predictions:
                class_counts = {}
                for pred in predictions.values():
                    cls = pred.get('class_friendly', pred.get('predicted_class', 'N/A'))
                    class_counts[cls] = class_counts.get(cls, 0) + 1
                
                consensus_list = [{'Clase': cls, 'Votos': count} for cls, count in class_counts.items()]
                df_consensus = pd.DataFrame(consensus_list)
                df_consensus.to_excel(writer, sheet_name='Consenso', index=False)
            
            # Formatear las hojas
            workbook = writer.book
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                worksheet.set_column('A:Z', 20)
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        logger.error(f"Error en generate_excel_report: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return None
